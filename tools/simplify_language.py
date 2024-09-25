import streamlit as st

import re
from datetime import datetime
from openai import OpenAI
import base64
from docx import Document
from docx.shared import Pt, Inches
import io
import numpy as np
from tools.utils_understandability import get_zix, get_cefr
import os
from tools.tool_base import (
    ToolBase,
    LOGFILE,
)


from tools.simplify_language_utils import (
    SYSTEM_MESSAGE_ES,
    SYSTEM_MESSAGE_LS,
    RULES_ES,
    RULES_LS,
    REWRITE_COMPLETE,
    REWRITE_CONDENSED,
    OPENAI_TEMPLATE_ES,
    OPENAI_TEMPLATE_LS,
    OPENAI_TEMPLATE_ANALYSIS_ES,
    OPENAI_TEMPLATE_ANALYSIS_LS,
    DEMO_TEXT,
)
from enum import Enum

from helper import init_logging, extract_text_from_uploaded_file, get_var

logger = init_logging(__name__, LOGFILE)

FILE_FORMAT_OPTIONS = ["pdf", "txt"]
INPUT_FORMAT_OPTIONS = ["Demo", "Datei Hochladen", "Text"]

# Height of the text areas for input and output.
TEXT_AREA_HEIGHT = 600

# Maximum number of characters for the input text.
# This is way below the context window sizes of the models.
# Adjust to your needs. However, we found that users can work and validate better when we nudge to work with shorter texts.
MAX_CHARS_INPUT = 10_000

# From our testing we derive a sensible temperature of 0.5 as a good trade-off between creativity and coherence. Adjust this to your needs.
TEMPERATURE = 0.5

# Constants for the formatting of the Word document that can be downloaded.
FONT_WORDDOC = "Arial"
FONT_SIZE_HEADING = 12
FONT_SIZE_PARAGRAPH = 9
FONT_SIZE_FOOTER = 7

# Limits for the understandability score to determine if the text is easy, medium or hard to understand.
LIMIT_HARD = 0
LIMIT_MEDIUM = -2


class InputFormat(Enum):
    DEMO = 0
    FILE = 1
    TEXT = 2


class SimplifyLanguage(ToolBase):
    def __init__(self, logger):
        super().__init__(logger)
        self.logger = logger
        self.title = "Sprache vereinfachen"
        self.model = "gpt-4o-mini"
        self.script_name, script_extension = os.path.splitext(__file__)

        self.input_files = []
        self.output_file = None
        self.results = []
        self.input_prefix = "input/"
        self.output_prefix = "output/"
        self.text = ""

        st.session_state.method = "Leichte Sprache"
        self.condense_text = False

        self.intro = self.get_intro()
        self.input_format = INPUT_FORMAT_OPTIONS[0]

        self.openai_client = OpenAI(
            api_key=get_var("OPENAI_API_KEY"),
        )

    def show_settings(self):
        st.markdown("Einstellungen")

        cols = st.columns([1, 1, 2])
        with cols[0]:
            self.method = st.radio(
                "Verarbeitungsmethode",
                [
                    "Leichte Sprache",
                    "Einfache Sprache",
                    "Analyse Leichte Sprache",
                    "Analyse Einfache Sprache",
                ],
                help="Siehe Information für Details der Unterschiede.",
                key="method",
            )
            if st.session_state.method == "Leichte Sprache":
                self.condense_text = st.toggle(
                    "Text verdichten",
                    value=True,
                    help="**Schalter aktiviert**: Modell konzentriert sich auf essentielle Informationen und versucht, Unwichtiges wegzulassen. **Schalter nicht aktiviert**: Modell versucht, alle Informationen zu übernehmen.",
                )

        with cols[1]:
            self.input_format = st.radio(
                label="Input Format", options=INPUT_FORMAT_OPTIONS
            )

        st.markdown("---")
        if INPUT_FORMAT_OPTIONS.index(self.input_format) == InputFormat.FILE.value:
            self.input_file = st.file_uploader(
                "PDF oder Text Datei",
                type=FILE_FORMAT_OPTIONS,
                help="Laden Sie die Datei hoch, die Sie vereinfachen oder analysieren möchten.",
            )
        elif INPUT_FORMAT_OPTIONS.index(self.input_format) == InputFormat.TEXT.value:
            self.input_text_box = st.text_area(
                "Ausgangstext, den du vereinfachen oder analysieren möchtest",
                value=None,
                height=TEXT_AREA_HEIGHT,
                max_chars=MAX_CHARS_INPUT,
                key="key_textinput",
            )

    def create_prompt(self):
        """Create prompt and system message according the app settings."""
        if st.session_state.method == "Leichte Sprache":
            if self.condense_text:
                final_prompt = OPENAI_TEMPLATE_LS.format(
                    rules=RULES_LS, completeness=REWRITE_CONDENSED, prompt=self.text
                )
            else:
                final_prompt = OPENAI_TEMPLATE_LS.format(
                    rules=RULES_LS, completeness=REWRITE_COMPLETE, prompt=self.text
                )
            system = SYSTEM_MESSAGE_LS
        elif st.session_state.method == "Einfache Sprache":
            final_prompt = OPENAI_TEMPLATE_ES.format(
                rules=RULES_ES, completeness=REWRITE_COMPLETE, prompt=self.text
            )
            system = SYSTEM_MESSAGE_ES
        elif st.session_state.method == "Analyse Leichte Sprache":
            final_prompt = OPENAI_TEMPLATE_ANALYSIS_LS.format(
                rules=RULES_LS, prompt=self.text
            )
            system = SYSTEM_MESSAGE_LS
        elif st.session_state.method == "Analyse Einfache Sprache":
            final_prompt = OPENAI_TEMPLATE_ANALYSIS_ES.format(
                rules=RULES_ES, prompt=self.text
            )
            system = SYSTEM_MESSAGE_ES

        return final_prompt, system

    def get_result_from_response(self, response):
        """Extract text between tags from response."""
        if "Leichte Sprache" in st.session_state.method:
            result = re.findall(
                r"<leichtesprache>(.*?)</leichtesprache>", response, re.DOTALL
            )
        else:
            result = re.findall(
                r"<einfachesprache>(.*?)</einfachesprache>", response, re.DOTALL
            )
        result = "\n".join(result)
        return result.strip()

    def invoke_openai_model(
        self,
        temperature=TEMPERATURE,
        max_tokens=4096,
    ):
        """Invoke OpenAI model."""
        final_prompt, system = self.create_prompt()
        try:
            message = self.openai_client.chat.completions.create(
                model=self.model,
                temperature=temperature,
                max_tokens=max_tokens,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": final_prompt},
                ],
            )
            message = message.choices[0].message.content.strip()
            return True, self.get_result_from_response(message)
        except Exception as e:
            print(f"Error: {e}")
            return False, e

    def run(self):
        def get_input_text():
            if INPUT_FORMAT_OPTIONS.index(self.input_format) == InputFormat.FILE.value:
                if self.input_file:
                    self.text = extract_text_from_uploaded_file(self.input_file)
            elif (
                INPUT_FORMAT_OPTIONS.index(self.input_format) == InputFormat.TEXT.value
            ):
                if self.input_text_box:
                    self.text = self.input_text_box
            else:
                self.text = DEMO_TEXT

        get_input_text()
        # Instantiate empty containers for the text areas.
        cols = st.columns([2, 2, 1])

        with cols[0]:
            source_text = st.container()
        with cols[1]:
            placeholder_result = st.empty()
        with cols[2]:
            placeholder_analysis = st.empty()

        # Populate containers.
        with source_text:
            st.text_area(
                "Ausgangstext, den du vereinfachen möchtest",
                value=self.text,
                height=TEXT_AREA_HEIGHT,
                max_chars=MAX_CHARS_INPUT,
                key="source_text",
            )
        button_action = (
            "Analysieren"
            if st.session_state.method.startswith("Analyse")
            else "Vereinfachen"
        )
        if st.button(button_action):
            with placeholder_result:
                text_output = st.text_area(
                    "Ergebnis",
                    height=TEXT_AREA_HEIGHT,
                )
            with placeholder_analysis:
                text_analysis = st.metric(
                    label="Verständlichkeit -10 bis 10",
                    value=None,
                    delta=None,
                    help="Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher, Texte in Leichter Sprache 2 bis 6 oder höher",
                )

            score_source = get_zix(self.text)
            # We add 0 to avoid negative zero.
            score_source_rounded = int(np.round(score_source, 0) + 0)
            cefr_source = get_cefr(score_source)

            # Analyze source text and display results.
            with source_text:
                if score_source < LIMIT_HARD:
                    st.markdown(
                        f"Dein Ausgangstext ist **:red[schwer verständlich]**. ({score_source_rounded} auf einer Skala von -10 bis 10). Das entspricht etwa dem **:red[Sprachniveau {cefr_source}]**."
                    )
                elif score_source >= LIMIT_HARD and score_source < LIMIT_MEDIUM:
                    st.markdown(
                        f"Dein Ausgangstext ist **:orange[nur mässig verständlich]**. ({score_source_rounded} auf einer Skala von -10 bis 10). Das entspricht etwa dem **:orange[Sprachniveau {cefr_source}]**."
                    )
                else:
                    st.markdown(
                        f"Dein Ausgangstext ist **:green[gut verständlich]**. ({score_source_rounded} auf einer Skala von -10 bis 10). Das entspricht etwa dem **:green[Sprachniveau {cefr_source}]**."
                    )
                with placeholder_analysis.container():
                    text_analysis = st.metric(
                        label="Verständlichkeit von -10 bis 10",
                        value=score_source_rounded,
                        delta=None,
                        help="Verständlichkeit auf einer Skala von -10 bis 10 Punkten (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher, Texte in Leichter Sprache 2 bis 6 oder höher.",
                    )

                with placeholder_analysis.container():
                    with st.spinner("Ich arbeite..."):
                        # Regular text simplification or analysis.
                        success, response = self.invoke_openai_model()

                if success is False:
                    st.error(
                        "Es ist ein Fehler bei der Abfrage der APIs aufgetreten. Bitte versuche es erneut. Alternativ überprüfe Code, API-Keys, Verfügbarkeit der Modelle und ggf. Internetverbindung."
                    )

                # Display results in UI.
                text = "Dein vereinfachter Text"
                if st.session_state.method.startswith("Analyse"):
                    text = "Deine Analyse"
                # Often the models return the German letter «ß». Replace it with the Swiss «ss».
                response = response.replace("ß", "ss")

                with placeholder_result.container():
                    st.text_area(
                        text,
                        height=TEXT_AREA_HEIGHT,
                        value=response,
                    )
                    if not st.session_state.method.startswith("Analyse"):
                        score_target = get_zix(response)
                        score_target_rounded = int(np.round(score_target, 0) + 0)
                        cefr_target = get_cefr(score_target)
                        if score_target < LIMIT_HARD:
                            st.markdown(
                                f"Dein vereinfachter Text ist **:red[schwer verständlich]**. ({score_target_rounded}  auf einer Skala von -10 bis 10). Das entspricht etwa dem **:red[Sprachniveau {cefr_target}]**."
                            )
                        elif score_target >= LIMIT_HARD and score_target < LIMIT_MEDIUM:
                            st.markdown(
                                f"Dein vereinfachter Text ist **:orange[nur mässig verständlich]**. ({score_target_rounded}  auf einer Skala von -10 bis 10). Das entspricht etwa dem **:orange[Sprachniveau {cefr_target}]**."
                            )
                        else:
                            st.markdown(
                                f"Dein vereinfachter Text ist **:green[gut verständlich]**. ({score_target_rounded}  auf einer Skala von -10 bis 10). Das entspricht etwa dem **:green[Sprachniveau {cefr_target}]**."
                            )
                        with placeholder_analysis.container():
                            text_analysis = st.metric(
                                label="Verständlichkeit -10 bis 10",
                                value=score_target_rounded,
                                delta=int(np.round(score_target - score_source, 0)),
                                help="Verständlichkeit auf einer Skala von -10 bis 10 (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher.",
                            )

                            create_download_link(self.text, response)
                    else:
                        with placeholder_analysis.container():
                            text_analysis = st.metric(
                                label="Verständlichkeit -10 bis 10",
                                value=score_source_rounded,
                                help="Verständlichkeit auf einer Skala von -10 bis 10 (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher.",
                            )
                            create_download_link(self.text, response, analysis=True)


def create_download_link(text_input, response, analysis=False):
    """Create a downloadable Word document and download link of the results."""
    document = Document()
    h1 = document.add_heading("Ausgangstext")
    p1 = document.add_paragraph("\n" + text_input)
    if analysis:
        h2 = document.add_heading("Analyse von Sprachmodell")
    else:
        h2 = document.add_heading("Vereinfachter Text von Sprachmodell")
    p2 = document.add_paragraph(response)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer = document.sections[0].footer
    footer.paragraphs[
        0
    ].text = f"Erstellt am {timestamp} mit der Prototyp-App «Einfache Sprache», Statistisches Amt, Kanton Basel-Stadt.\n"
    # Set font for all paragraphs.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_WORDDOC
    # Set font size for all headings.
    for paragraph in [h1, h2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_HEADING)
    # Set font size for all paragraphs.
    for paragraph in [p1, p2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_PARAGRAPH)
    # Set font and font size for footer.
    for run in footer.paragraphs[0].runs:
        run.font.name = "Arial"
        run.font.size = Pt(FONT_SIZE_FOOTER)
    section = document.sections[0]
    section.page_width = Inches(8.27)  # Width of A4 paper in inches
    section.page_height = Inches(11.69)  # Height of A4 paper in inches
    io_stream = io.BytesIO()
    document.save(io_stream)
    # # A download button unfortunately resets the app. So we use a link instead.
    # https://github.com/streamlit/streamlit/issues/4382#issuecomment-1223924851
    # https://discuss.streamlit.io/t/creating-a-pdf-file-generator/7613?u=volodymyr_holomb
    b64 = base64.b64encode(io_stream.getvalue())
    file_name = "Ergebnis.docx"
    caption = "Vereinfachten Text herunterladen"
    if analysis:
        file_name = "Analyse.docx"
        caption = "Analyse herunterladen"
    download_url = f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{file_name}">{caption}</a>'
    st.markdown(download_url, unsafe_allow_html=True)
